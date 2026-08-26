#!/usr/bin/env python3
"""Generate the end-user guide for FundingPlatform as an editable DOCX."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "user-guide" / "Manual-de-Usuario-FundingPlatform.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
PALE_GREEN = "EDF7F0"
PALE_AMBER = "FFF7E6"
PALE_RED = "FDEEEE"
MID_GRAY = "667085"
LIGHT_GRAY = "F4F5F7"
WHITE = "FFFFFF"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D0D5DD", size=6, sides=("top", "left", "bottom", "right")) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in sides:
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), color)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_repeat_table_header_and_keep(table) -> None:
    if table.rows:
        repeat_table_header(table.rows[0])
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def set_font(run, name="Calibri", size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.add_run("Página ")
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Small note" not in styles:
        note = styles.add_style("Small note", WD_STYLE_TYPE.PARAGRAPH)
        note.base_style = normal
        note.font.name = "Calibri"
        note.font.size = Pt(9)
        note.font.color.rgb = RGBColor.from_string(MID_GRAY)
        note.paragraph_format.space_after = Pt(4)
        note.paragraph_format.line_spacing = 1.15


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(1.7)
    repeat_table_header(table.rows[0])
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(left.add_run("FUNDINGPLATFORM"), size=9, bold=True, color=NAVY)
    set_font(right.add_run("MANUAL DE USUARIO"), size=8, bold=True, color=MID_GRAY)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 0, 0, 60, 0)
        set_cell_border(cell, color=BLUE, size=8, sides=("bottom",))

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer_p.add_run("FundingPlatform · Guía para organizaciones · "), size=8, color=MID_GRAY)
    add_page_number(footer_p)
    for run in footer_p.runs:
        set_font(run, size=8, color=MID_GRAY)


def add_title(doc: Document, text: str, subtitle: str) -> None:
    hero = doc.add_table(rows=1, cols=1)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    hero.autofit = False
    hero.columns[0].width = Inches(6.5)
    repeat_table_header(hero.rows[0])
    cell = hero.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, NAVY)
    set_cell_margins(cell, 700, 360, 700, 360)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("RISE FUNDING\n")
    set_font(r, size=12, bold=True, color="B9D7F0")
    r = p.add_run(text)
    set_font(r, size=28, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(12)
    r = p2.add_run(subtitle)
    set_font(r, size=14, color=WHITE)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(intro.add_run("Guía práctica para equipos de organizaciones sin fines de lucro"), size=14, bold=True, color=DARK_BLUE)
    meta = doc.add_paragraph()
    set_font(meta.add_run("Versión 1.0 · Agosto de 2026\n"), size=10, bold=True, color=MID_GRAY)
    set_font(meta.add_run("Ámbito: usuarios de organizaciones. La consola administrativa se documenta por separado."), size=10, color=MID_GRAY)

    callout(doc, "Propósito", "Aprender a preparar el perfil institucional y los proyectos, descubrir fondos, interpretar la compatibilidad, organizar postulaciones y mantener un calendario común.", "info")
    callout(doc, "Dos límites importantes", "FundingPlatform no envía postulaciones al financiador y sus resultados de compatibilidad no confirman elegibilidad. Siempre debes revisar las bases y la fuente oficial.", "warning")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(item))


def add_steps(doc: Document, items: list[str]) -> None:
    numbering = doc.part.numbering_part.element
    source_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    source_num = next(node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == source_id)
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(new_num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id)
        set_font(p.add_run(item))


def callout(doc: Document, title: str, text: str, kind="info") -> None:
    colors = {"info": PALE_BLUE, "success": PALE_GREEN, "warning": PALE_AMBER, "danger": PALE_RED}
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_fill(cell, colors[kind])
    set_cell_border(cell, color=BLUE if kind == "info" else "C68A00", size=8, sides=("left",))
    set_cell_margins(cell, 120, 180, 120, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(title), size=10, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run(text), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def screenshot_slot(doc: Document, number: int, title: str, what_to_show: str) -> None:
    visual_steps = {
        1: ["Correo", "Contraseña", "Iniciar sesión", "Recuperar contraseña", "Microsoft, si está habilitado"],
        2: ["Identidad", "Impacto", "Financiamiento", "Porcentaje de avance", "Guardar"],
        3: ["Estado y versión", "Título y resumen", "Brecha de financiamiento", "Ver o Editar", "Calcular compatibilidad"],
        4: ["Texto y filtros", "Ordenar", "Resultados", "Guardar búsqueda", "Abrir oportunidad"],
        5: ["Seleccionar proyecto", "Calcular", "Compatible / Incompatible / Datos insuficientes", "Cobertura y razones"],
        6: ["Proyecto + oportunidad", "Estado y fechas", "Guardar postulación", "Mes anterior / siguiente", "Abrir hito"],
    }
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_fill(cell, LIGHT_GRAY)
    set_cell_border(cell, color="98A2B3", size=6)
    set_cell_margins(cell, 260, 240, 260, 240)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(f"GUÍA VISUAL {number} · {title}"), size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p2.add_run("  →  ".join(visual_steps[number])), size=9, bold=True, color=NAVY)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p3.add_run(what_to_show), size=9, color=MID_GRAY, italic=True)
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(6)
    p4.paragraph_format.space_after = Pt(0)
    set_font(p4.add_run(f"Figura {number}. Esquema orientativo de {title.lower()}; la posición exacta puede variar según pantalla y versión."), size=8, italic=True, color=MID_GRAY)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, LIGHT_BLUE)
        set_cell_border(cell)
        set_cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(header), size=10, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
            set_font(cells[idx].paragraphs[0].add_run(value), size=9.5)
    set_repeat_table_header_and_keep(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "Manual de usuario de FundingPlatform"
    doc.core_properties.subject = "Guía práctica para organizaciones usuarias de Rise Funding"
    doc.core_properties.author = "Rise Funding"
    doc.core_properties.keywords = "fondos concursables, ONG, proyectos, compatibilidad, postulaciones"

    add_title(doc, "Manual de usuario", "Cómo organizar tus proyectos y oportunidades de financiamiento")
    page_break(doc)

    add_heading(doc, "Contenido", 1)
    add_table(doc, ["Sección", "Qué aprenderás"], [
        ["1. Antes de comenzar", "Roles, disponibilidad y conceptos básicos."],
        ["2. Acceso y primera configuración", "Inicio de sesión, organización y perfil institucional."],
        ["3. Navegación", "Qué encontrarás en cada módulo."],
        ["4. Proyectos", "Crear, completar, publicar y archivar proyectos."],
        ["5. Concursos, favoritos y alertas", "Buscar, filtrar, guardar y revisar oportunidades."],
        ["6. Compatibilidad", "Calcular e interpretar resultados orientativos."],
        ["7. Postulaciones y calendario", "Organizar el trabajo interno y sus fechas."],
        ["8. Conexiones y marketplace", "Visibilidad pública y red de organizaciones."],
        ["9. Cuenta, seguridad y suscripción", "Acceso, MFA, Microsoft y estado del plan."],
        ["10. Flujos recomendados", "Recorridos completos para tareas habituales."],
        ["11. Solución de problemas", "Qué revisar y cuándo pedir ayuda."],
        ["12. Glosario", "Significado de los términos principales."],
    ], [1.9, 4.6])
    callout(doc, "Sobre las guías visuales", "Esta edición incluye seis esquemas orientativos basados en la interfaz real. Los pantallazos del ambiente dev se podrán incorporar después del despliegue, usando únicamente datos ficticios y sin exponer información personal.", "info")
    add_body(doc, "Recomendación: usa la versión PDF para compartir y la versión Word para mantener el manual actualizado cuando cambie la interfaz.")
    page_break(doc)

    add_heading(doc, "1. Antes de comenzar", 1)
    add_heading(doc, "1.1 ¿Qué permite hacer FundingPlatform?", 2)
    add_body(doc, "FundingPlatform ayuda a una organización a ordenar la información institucional y de sus proyectos, descubrir concursos, revisar compatibilidad estructurada y llevar un seguimiento interno de postulaciones y fechas.")
    add_bullets(doc, [
        "Crear y mantener el perfil de la organización.",
        "Registrar proyectos con su impacto, presupuesto y brecha de financiamiento.",
        "Buscar concursos mediante texto, filtros, fechas, montos y categorías.",
        "Guardar favoritos y búsquedas privadas.",
        "Calcular compatibilidad orientativa por proyecto y conservar el historial.",
        "Registrar postulaciones y visualizar hitos en un calendario común.",
        "Participar, si la organización lo autoriza, en un directorio de conexiones.",
    ])
    callout(doc, "La plataforma no sustituye el trabajo de postulación", "No completa ni envía formularios al financiador, no garantiza adjudicación y no reemplaza la lectura de las bases oficiales.", "warning")

    add_heading(doc, "1.2 Roles", 2)
    add_table(doc, ["Rol", "Puede hacer", "Consideraciones"], [
        ["Visitante", "Explorar fondos y proyectos públicos.", "No ve perfiles internos, favoritos ni postulaciones."],
        ["Miembro de organización", "Usar el espacio privado según sus permisos.", "La edición puede depender del rol asignado dentro de la organización."],
        ["Administrador de organización", "Gestionar visibilidad y solicitudes de conexión.", "Debe revisar cuidadosamente los datos que se vuelven públicos."],
        ["Administrador de plataforma", "Operar catálogo, usuarios y procesos internos.", "No está cubierto en detalle por este manual."],
    ], [1.4, 2.7, 2.4])

    add_heading(doc, "1.3 Funciones sujetas al ambiente", 2)
    add_body(doc, "Durante un piloto algunas integraciones pueden estar temporalmente deshabilitadas. La interfaz mostrará un mensaje si una función no está disponible.")
    add_table(doc, ["Función", "Qué esperar"], [
        ["Registro y recuperación por correo", "Puede estar deshabilitado en el ambiente dev. El equipo administrador entregará o restaurará el acceso."],
        ["Inicio de sesión con Microsoft", "Aparece cuando la integración con Microsoft Entra está configurada."],
        ["Alertas por correo", "La búsqueda puede quedar guardada aunque el envío todavía no esté habilitado."],
        ["Planes pagados", "Los cobros están deshabilitados o restringidos a sandbox en esta etapa."],
    ], [2.3, 4.2])
    add_heading(doc, "2. Acceso y primera configuración", 1)
    add_heading(doc, "2.1 Iniciar sesión", 2)
    add_steps(doc, [
        "Abre la dirección entregada por el equipo de implementación.",
        "Selecciona Iniciar sesión e ingresa tu correo y contraseña.",
        "Si aparece un desafío MFA, escribe el código de tu aplicación autenticadora o utiliza un código de recuperación.",
        "Al ingresar por primera vez, la plataforma te enviará al onboarding si todavía no perteneces a una organización.",
    ])
    callout(doc, "Sesión segura", "Si te alejas del equipo, bloquea la pantalla. Al terminar, usa el icono Cerrar sesión; cerrar solo la pestaña no reemplaza esta acción.", "info")
    screenshot_slot(doc, 1, "Acceso a la plataforma", "El recorrido comienza con la cuenta local y puede incorporar Microsoft o MFA según la configuración.")

    add_heading(doc, "2.2 Crear el espacio de la organización", 2)
    add_steps(doc, [
        "En el onboarding, escribe el nombre público de la organización.",
        "Selecciona el país principal y el tipo de organización.",
        "Selecciona Crear organización.",
        "Confirma que aparece el perfil institucional con su porcentaje de avance.",
    ])
    add_body(doc, "El espacio separa los datos, proyectos y resultados de tu organización de los de otras organizaciones.")

    add_heading(doc, "2.3 Completar el perfil institucional", 2)
    add_body(doc, "El perfil se organiza en tres pasos: Identidad, Impacto y Financiamiento. Guarda los cambios antes de abandonar la pantalla.")
    add_table(doc, ["Paso", "Información principal", "Por qué importa"], [
        ["Identidad", "Nombre, razón social, identificador tributario, año, tipo, personalidad jurídica, tamaño, web y descripción.", "Sustenta condiciones legales y antigüedad."],
        ["Impacto", "Países, regiones, áreas, beneficiarios y tipos de proyecto.", "Permite comparar territorios y taxonomías."],
        ["Financiamiento", "Experiencia, monedas, rangos de financiamiento, presupuesto, idiomas y resumen.", "Mejora la calidad de los datos usados en búsquedas y compatibilidad."],
    ], [1.2, 3.1, 2.2])
    callout(doc, "Calidad de datos", "Un dato desconocido no se considera aprobado. Completa la información con evidencia interna y actualízala cuando cambie.", "warning")
    screenshot_slot(doc, 2, "Perfil de la organización", "El avance reúne los tres grupos de datos; cada grupo se revisa y guarda desde el mismo perfil.")
    add_heading(doc, "3. Navegación del espacio privado", 1)
    add_body(doc, "En escritorio, el menú aparece a la izquierda. En pantallas pequeñas, se muestra como una franja horizontal desplazable. La barra superior permite cambiar el tema y cerrar la sesión.")
    add_table(doc, ["Menú", "Uso principal"], [
        ["Resumen", "Punto de entrada al espacio de la organización."],
        ["Concursos disponibles", "Buscar y filtrar oportunidades vigentes."],
        ["Compatibilidad", "Comparar un proyecto con fondos activos y revisar el historial."],
        ["Favoritos", "Reencontrar concursos guardados por la cuenta en la organización."],
        ["Postulaciones", "Registrar el avance interno de cada postulación."],
        ["Calendario", "Revisar cierres e hitos mensuales."],
        ["Alertas", "Guardar búsquedas y, cuando esté disponible, recibir resúmenes."],
        ["Conexiones", "Explorar el directorio y administrar invitaciones moderadas."],
        ["Organización", "Editar el perfil institucional."],
        ["Proyectos", "Crear y mantener proyectos."],
        ["Mi cuenta", "Revisar métodos de acceso y vincular Microsoft cuando esté habilitado."],
        ["Suscripción", "Consultar el plan y el estado del acceso de la organización."],
    ], [2.0, 4.5])
    callout(doc, "Si un menú no aparece", "Puede deberse a tu rol, al estado de la organización o a que el módulo todavía no está habilitado en ese ambiente.", "info")
    add_heading(doc, "4. Gestionar proyectos", 1)
    add_heading(doc, "4.1 Crear un proyecto", 2)
    add_steps(doc, [
        "Entra a Proyectos y selecciona Nuevo proyecto.",
        "Completa el título, resumen, descripción, estado, territorios, categorías, beneficiarios y tipo de proyecto.",
        "Registra fechas y presupuesto. Define la moneda de forma consistente.",
        "Revisa el financiamiento confirmado y la brecha pendiente.",
        "Guarda y vuelve a abrir el proyecto para confirmar la información.",
    ])
    add_body(doc, "Cada proyecto mantiene su propia versión, necesidad de financiamiento e historial. Esto evita mezclar proyectos con territorios, poblaciones o presupuestos diferentes.")

    add_heading(doc, "4.2 Estados y publicación", 2)
    add_table(doc, ["Concepto", "Qué significa"], [
        ["Estado del proyecto", "Describe su avance: idea, diseño, buscando financiamiento, financiado parcialmente, financiado, en ejecución o completado."],
        ["Estado de publicación", "Controla si el proyecto está en borrador, revisión, publicado, rechazado o archivado."],
        ["Versión", "Aumenta cuando cambia información relevante y permite detectar resultados de compatibilidad antiguos."],
    ], [2.0, 4.5])
    add_bullets(doc, [
        "Para solicitar publicación, completa los requisitos mostrados en el panel del proyecto.",
        "El perfil de la organización debe estar completo al menos en el porcentaje exigido por la interfaz.",
        "Un proyecto archivado conserva historial y postulaciones, pero no admite cálculos nuevos de compatibilidad.",
    ])
    screenshot_slot(doc, 3, "Lista y detalle de proyectos", "La tarjeta resume el proyecto y el detalle permite editarlo o iniciar un cálculo vigente.")

    add_heading(doc, "4.3 Buenas prácticas", 2)
    add_bullets(doc, [
        "Usa un título específico y un resumen que pueda entender una persona externa.",
        "No incluyas contraseñas, datos bancarios ni información personal innecesaria.",
        "Actualiza territorios, beneficiarios, fechas y presupuesto antes de calcular compatibilidad.",
        "Si cambió el proyecto, crea un cálculo nuevo; no tomes una ejecución antigua como vigente.",
    ])
    add_heading(doc, "5. Concursos, favoritos y alertas", 1)
    add_heading(doc, "5.1 Buscar oportunidades", 2)
    add_steps(doc, [
        "Abre Concursos disponibles.",
        "Escribe palabras relevantes o aplica filtros de patrocinador, monto, moneda, fecha, territorio, categoría, beneficiarios, tipo de proyecto u organización.",
        "Ordena por relevancia, cierre próximo, publicación reciente o monto máximo.",
        "Abre el detalle del concurso y revisa requisitos, montos, fechas, patrocinador y fuente oficial.",
        "Si el concurso interesa, guárdalo en Favoritos o inicia una postulación interna.",
    ])
    callout(doc, "Fecha oficial", "El calendario y las fichas ayudan a organizarse, pero siempre confirma la fecha y hora en la fuente oficial. Una convocatoria continua se muestra como tal y no debe confundirse con una fecha desconocida.", "warning")
    screenshot_slot(doc, 4, "Catálogo de concursos", "Los filtros modifican los resultados; cada tarjeta conduce a la ficha y a su fuente oficial.")

    add_heading(doc, "5.2 Favoritos", 2)
    add_body(doc, "El botón de favorito permite guardar o quitar una oportunidad. La sección Favoritos reúne los concursos guardados por tu cuenta dentro de la organización.")
    add_bullets(doc, [
        "Guardar un concurso no crea una postulación.",
        "Quitar un favorito no elimina una postulación existente.",
        "Los favoritos pueden aportar fechas al calendario cuando la oportunidad tiene cierre conocido.",
    ])

    add_heading(doc, "5.3 Búsquedas y alertas", 2)
    add_steps(doc, [
        "Aplica los filtros deseados en Concursos disponibles.",
        "Selecciona Guardar búsqueda.",
        "Asigna un nombre descriptivo, por ejemplo: Fondos ambientales Chile 2026.",
        "Activa el resumen diario y elige una hora local si el correo está habilitado.",
        "Desde Alertas puedes abrir, activar, desactivar o eliminar cada búsqueda.",
    ])
    callout(doc, "Privacidad", "Las búsquedas guardadas son privadas. Una alerta informa oportunidades nuevas, pero no confirma que tu organización pueda postular.", "info")
    add_heading(doc, "6. Interpretar la compatibilidad", 1)
    add_heading(doc, "6.1 Calcular", 2)
    add_steps(doc, [
        "Actualiza el perfil institucional y el proyecto.",
        "Entra a Compatibilidad y selecciona un proyecto.",
        "Selecciona Calcular compatibilidad.",
        "Espera a que aparezca la ejecución completada y revisa sus resultados.",
        "Abre cada fondo para leer condiciones excluyentes, razones y datos utilizados.",
    ])
    add_body(doc, "El cálculo compara datos estructurados de un proyecto con un máximo acotado de fondos activos. Es determinístico y reproducible: no usa IA para decidir ni cambia los requisitos del financiador.")

    add_heading(doc, "6.2 Estados", 2)
    add_table(doc, ["Resultado", "Interpretación correcta", "Acción recomendada"], [
        ["Compatible", "No se detectaron fallos en las condiciones excluyentes evaluadas y existen datos suficientes.", "Revisar igualmente las bases completas y preparar una postulación."],
        ["Incompatible", "Al menos una condición excluyente evaluada no se cumple.", "Abrir las razones, verificar el dato y descartar o corregir si corresponde."],
        ["Datos insuficientes", "Falta información para decidir una o más condiciones excluyentes.", "Completar el perfil/proyecto o revisar manualmente el requisito."],
    ], [1.35, 3.15, 2.0])
    callout(doc, "Desconocido nunca equivale a aprobado", "Los datos ausentes aportan cero al puntaje de esa regla y reducen la cobertura. No se redistribuyen puntos para ocultar información faltante.", "warning")

    add_heading(doc, "6.3 Puntaje, cobertura y vigencia", 2)
    add_bullets(doc, [
        "Las condiciones excluyentes se muestran separadas del puntaje.",
        "Si una condición excluyente falla, el resultado es Incompatible y el puntaje puede mostrarse como No aplica.",
        "La cobertura indica cuánto de la comparación pudo evaluarse con datos disponibles.",
        "Las razones explican qué regla pasó, falló o quedó desconocida; revisa la evidencia mostrada.",
        "Una ejecución deja de estar vigente si cambia el proyecto, el perfil, el contenido del fondo o la configuración de reglas.",
        "Los proyectos archivados permiten consultar el historial, pero no calcular una versión nueva.",
    ])
    screenshot_slot(doc, 5, "Resultados de compatibilidad", "El estado global, la cobertura y las razones se leen juntos; ninguno confirma elegibilidad.")
    callout(doc, "Aviso que siempre debes recordar", "Resultado orientativo basado en datos disponibles; no confirma elegibilidad ni reemplaza la revisión de las bases del fondo.", "danger")
    add_heading(doc, "7. Postulaciones y calendario", 1)
    add_heading(doc, "7.1 Iniciar y mantener una postulación", 2)
    add_body(doc, "Una postulación es un registro interno que relaciona un proyecto con una oportunidad. No se envía al financiador desde FundingPlatform.")
    add_steps(doc, [
        "Desde el detalle de un concurso, selecciona Iniciar postulación; también puedes usar Nueva postulación en el módulo Postulaciones.",
        "Selecciona el proyecto correspondiente y la oportunidad.",
        "Registra notas internas, fechas planificadas, monto solicitado y moneda cuando corresponda.",
        "Actualiza el estado a medida que avanza el trabajo.",
        "Guarda y confirma el mensaje de éxito antes de cerrar la pantalla.",
    ])
    add_table(doc, ["Estado", "Uso sugerido"], [
        ["Interesada", "La organización decidió explorar la oportunidad."],
        ["Preparando", "El equipo está reuniendo antecedentes o redactando."],
        ["Presentada", "La postulación fue enviada fuera de FundingPlatform."],
        ["Adjudicada", "El financiador confirmó un resultado favorable."],
        ["Rechazada", "La postulación no fue seleccionada."],
        ["Descartada", "La organización decidió no continuar."],
    ], [1.6, 4.9])
    callout(doc, "Conflicto de edición", "Si otra persona guardó cambios mientras tenías la ficha abierta, la plataforma puede pedirte recargar la versión actual. Revisa la información antes de volver a guardar.", "info")

    add_heading(doc, "7.2 Usar el calendario", 2)
    add_body(doc, "El calendario reúne cierres e hitos registrados en proyectos, postulaciones y favoritos. Puedes avanzar o retroceder por mes y abrir el elemento relacionado.")
    add_bullets(doc, [
        "Mantén las fechas de proyecto y postulación actualizadas.",
        "No confíes solo en el calendario para una fecha de cierre; revisa la fuente oficial.",
        "Si no hay hitos en un mes, revisa otro mes o agrega fechas a proyectos y postulaciones.",
    ])
    screenshot_slot(doc, 6, "Postulaciones y calendario", "La postulación conserva su avance interno y las fechas relacionadas aparecen en la agenda mensual.")
    add_heading(doc, "8. Conexiones y marketplace", 1)
    add_heading(doc, "8.1 Marketplace público", 2)
    add_body(doc, "El marketplace muestra proyectos publicados y perfiles públicos seguros. No expone borradores ni datos privados de la organización.")
    add_bullets(doc, [
        "Verifica qué información será pública antes de solicitar la publicación de un proyecto.",
        "Un perfil público de organización solo aparece cuando cumple las condiciones de visibilidad.",
        "No publiques correos personales, teléfonos privados, identificadores tributarios ni notas internas en descripciones públicas.",
    ])

    add_heading(doc, "8.2 Directorio y conexiones", 2)
    add_body(doc, "La organización no aparece en el directorio por defecto. Un administrador de la organización puede activar la visibilidad y decidir si acepta nuevas solicitudes.")
    add_steps(doc, [
        "Entra a Conexiones y revisa Privacidad y visibilidad.",
        "Si tienes permisos, activa el directorio y la recepción de solicitudes.",
        "Busca organizaciones por nombre o descripción y abre su perfil público.",
        "Para invitar, elige un propósito, un proyecto público opcional y escribe un mensaje breve sin datos de contacto.",
        "Revisa Solicitudes y conexiones para aceptar, rechazar o retirar acciones disponibles.",
    ])
    callout(doc, "Sin chat ni publicación automática de contacto", "El primer mensaje es privado para las dos organizaciones. Aceptar una conexión no publica automáticamente correos, teléfonos ni enlaces.", "info")
    add_heading(doc, "9. Cuenta, seguridad y suscripción", 1)
    add_heading(doc, "9.1 Seguridad de la cuenta", 2)
    add_bullets(doc, [
        "Usa una contraseña única de al menos 12 caracteres cuando el acceso local esté habilitado.",
        "No compartas códigos MFA ni códigos de recuperación.",
        "Guarda los códigos de recuperación en un lugar seguro; se muestran una sola vez al activar MFA.",
        "Desde Mi cuenta puedes vincular una identidad Microsoft cuando la integración esté disponible.",
        "Una cuenta Microsoft personal y una laboral con el mismo correo son identidades distintas: selecciona la correcta.",
        "Cierra sesión al terminar, especialmente en equipos compartidos.",
    ])

    add_heading(doc, "9.2 Suscripción", 2)
    add_body(doc, "La suscripción pertenece a la organización. En la etapa actual, los cobros están deshabilitados o limitados a pruebas. La pantalla Suscripción informa el plan y su estado, pero no debes ingresar datos de tarjeta si el equipo de implementación no ha anunciado un flujo de pago habilitado.")
    callout(doc, "Ante una solicitud inesperada", "Si una pantalla o mensaje solicita datos bancarios fuera del flujo informado, detente y consulta al administrador de la plataforma.", "danger")
    add_heading(doc, "10. Flujos recomendados", 1)
    add_heading(doc, "10.1 Preparar la organización para buscar fondos", 2)
    add_steps(doc, [
        "Completa Identidad, Impacto y Financiamiento del perfil institucional.",
        "Crea un proyecto con territorios, beneficiarios, categorías, fechas y presupuesto.",
        "Revisa la brecha de financiamiento y guarda.",
        "Explora Concursos disponibles y guarda los más relevantes.",
        "Calcula compatibilidad y revisa razones, cobertura y vigencia.",
        "Lee las bases oficiales antes de iniciar una postulación.",
    ])

    add_heading(doc, "10.2 Organizar una postulación", 2)
    add_steps(doc, [
        "Abre la oportunidad y confirma el cierre en la fuente oficial.",
        "Inicia la postulación y vincula el proyecto correcto.",
        "Registra responsables fuera de la plataforma si tu equipo lo requiere; FundingPlatform no reemplaza tu herramienta de colaboración.",
        "Agrega hitos y fechas; revisa el calendario del mes.",
        "Actualiza el estado cuando se presenta y cuando llega el resultado.",
    ])

    add_heading(doc, "10.3 Revisión semanal", 2)
    add_bullets(doc, [
        "Revisar concursos con cierre próximo.",
        "Verificar favoritos y búsquedas guardadas.",
        "Actualizar postulaciones y fechas del calendario.",
        "Recalcular compatibilidad solo cuando cambien datos relevantes.",
        "Archivar proyectos que ya no están activos sin perder su historial.",
    ])
    add_heading(doc, "11. Solución de problemas", 1)
    add_table(doc, ["Situación", "Qué revisar", "Qué hacer"], [
        ["No puedo iniciar sesión", "Correo, contraseña, verificación, MFA y ambiente correcto.", "Reintenta una vez. Si persiste, solicita restauración de acceso al administrador."],
        ["Registro o recuperación no disponible", "El ambiente puede tener correo deshabilitado.", "Pide acceso o recuperación manual al equipo del piloto."],
        ["No veo una organización", "La cuenta puede no tener membresía o el onboarding no terminó.", "Completa el onboarding o pide que revisen tu membresía."],
        ["No puedo editar", "Rol, estado publicado/archivado y cambios de otra persona.", "Recarga la ficha y confirma tus permisos."],
        ["No hay resultados de compatibilidad", "Proyecto archivado, datos incompletos, conexión o ausencia de fondos candidatos.", "Completa datos, revisa el estado y reintenta una sola vez."],
        ["Resultado Datos insuficientes", "Perfil o proyecto carece de datos usados por una condición.", "Abre las razones y completa solo información verificada."],
        ["Una fecha parece distinta", "Zona horaria y precisión de la fecha oficial.", "Confirma día y hora en la fuente del financiador."],
        ["La alerta no envía correo", "La búsqueda puede existir con correo deshabilitado.", "Usa Abrir desde Alertas y consulta el estado del ambiente."],
        ["Cambios rechazados al guardar", "Otra persona pudo actualizar la ficha.", "Recarga, compara y vuelve a aplicar tus cambios."],
    ], [1.8, 2.35, 2.35])

    add_heading(doc, "11.1 Información útil al pedir soporte", 2)
    add_bullets(doc, [
        "Fecha y hora aproximada del problema.",
        "Pantalla y acción realizada, sin incluir contraseñas ni códigos MFA.",
        "Mensaje mostrado por la plataforma.",
        "Navegador y dispositivo utilizados.",
        "Identificador público del proyecto, concurso o postulación, si está visible.",
    ])
    callout(doc, "Nunca envíes a soporte", "Contraseñas, tokens, códigos MFA, códigos de recuperación, datos bancarios o capturas que expongan información personal innecesaria.", "danger")
    add_heading(doc, "12. Glosario", 1)
    add_table(doc, ["Término", "Definición"], [
        ["Brecha de financiamiento", "Diferencia pendiente entre el presupuesto del proyecto y el financiamiento confirmado."],
        ["Cobertura", "Porcentaje ponderado de reglas que pudo evaluarse con datos disponibles."],
        ["Condición excluyente", "Requisito cuyo incumplimiento puede hacer que el resultado sea Incompatible."],
        ["Datos insuficientes", "Estado que indica información desconocida para una condición relevante."],
        ["Elegibilidad", "Cumplimiento real de todas las bases del financiador. FundingPlatform no la confirma."],
        ["ETag o versión de edición", "Mecanismo que evita sobrescribir silenciosamente cambios hechos por otra persona."],
        ["Favorito", "Concurso guardado para acceder rápidamente; no equivale a postulación."],
        ["Marketplace", "Catálogo público de proyectos y organizaciones que autorizaron su publicación."],
        ["MFA", "Segundo factor de autenticación mediante código temporal o recuperación."],
        ["Postulación", "Registro interno de seguimiento; el envío al financiador ocurre fuera de FundingPlatform."],
        ["Proyecto archivado", "Proyecto que conserva historial pero ya no admite ciertas acciones nuevas."],
        ["Vigencia", "Indica si un cálculo aún corresponde a las versiones actuales de los datos."],
    ], [2.0, 4.5])

    add_heading(doc, "Cierre", 1)
    add_body(doc, "FundingPlatform es más útil cuando el equipo mantiene datos completos, separa claramente cada proyecto y usa los resultados como apoyo para revisar oportunidades, no como una decisión automática.")
    callout(doc, "Regla final", "Antes de postular: confirma la fuente oficial, lee las bases completas, verifica la fecha y documenta internamente la decisión.", "success")
    note = doc.add_paragraph(style="Small note")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(note.add_run("Fin del manual · Versión 1.0 · Agosto de 2026"), size=9, bold=True, color=MID_GRAY)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
